"""
文档解析器：支持 PDF、DOCX、TXT、Markdown、Excel、图片 等格式。
使用 PyMuPDF 解析 PDF，python-docx 解析 Word 文档，
openpyxl 解析 Excel，Pillow + OCR 解析图片。

同时提取文档元数据（文件哈希、大小等），用于版本管理。
"""

import re
import hashlib
import os
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field
from datetime import datetime, timezone

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from loguru import logger


@dataclass
class ParsedPage:
    """解析后的单页内容"""
    page_num: int
    text: str
    metadata: dict = field(default_factory=dict)


@dataclass
class DocumentMeta:
    """
    文档元数据 —— 用于版本管理与冲突解决

    每个文档由 doc_id 唯一标识（基于文件路径），
    通过 file_hash 检测内容变更，version 记录更新次数。
    """
    doc_id: str                     # 文档唯一标识
    file_name: str                  # 原始文件名
    file_path: str                  # 文件路径
    file_hash: str                  # 文件内容 SHA256 哈希
    file_size: int = 0              # 文件大小（字节）
    file_type: str = ""             # 文件类型：pdf / docx / xlsx / txt / md / csv / image 等
    version: int = 1                # 版本号，每次更新递增
    created_at: str = ""            # 首次索引时间（ISO 8601）
    updated_at: str = ""            # 最近更新时间（ISO 8601）
    operator: str = "system"        # 操作者
    source: str = "upload"          # 来源：upload / api / batch
    status: str = "active"          # 状态：active / indexing / archived / outdated
    tags: list[str] = field(default_factory=list)  # 自定义标签
    vector_count: int = 0           # 子块向量数量（索引完成后回填）


@dataclass
class ParsedDocument:
    """解析后的完整文档"""
    file_path: str
    title: str = ""
    pages: list[ParsedPage] = field(default_factory=list)
    raw_text: str = ""
    toc: list[dict] = field(default_factory=list)
    doc_meta: Optional[DocumentMeta] = None  # 文档级元数据

    @property
    def full_text(self) -> str:
        if self.raw_text:
            return self.raw_text
        return "\n".join(p.text for p in self.pages)


class DocumentParser:
    """多格式文档解析器"""

    def parse(self, file_path: str, operator: str = "system") -> ParsedDocument:
        """解析文档，根据扩展名选择解析器"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()

        if ext == ".pdf":
            doc = self._parse_pdf(file_path)
        elif ext == ".docx":
            doc = self._parse_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            doc = self._parse_excel(file_path)
        elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"):
            doc = self._parse_image(file_path)
        elif ext in (".txt", ".md"):
            doc = self._parse_text(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

        # 提取文档元数据
        doc.doc_meta = self._extract_meta(file_path, operator)
        return doc

    def _extract_meta(self, file_path: str, operator: str) -> DocumentMeta:
        """提取文档元数据：哈希、大小、时间等"""
        path = Path(file_path)
        stat = path.stat()

        # 计算文件 SHA256
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        file_hash = sha256.hexdigest()

        now = datetime.now(timezone.utc).isoformat()

        # doc_id 基于文件路径的稳定哈希
        doc_id = hashlib.md5(str(path.absolute()).encode()).hexdigest()[:12]

        # 从文件后缀推断类型
        ext = path.suffix.lower().lstrip(".")
        file_type = ext

        return DocumentMeta(
            doc_id=doc_id,
            file_name=path.name,
            file_path=str(path.absolute()),
            file_hash=file_hash,
            file_size=stat.st_size,
            file_type=file_type,
            version=1,
            created_at=now,
            updated_at=now,
            operator=operator,
            source="upload",
            status="indexing",  # 索引中，完成后更新为 active
            vector_count=0,
        )

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """解析 PDF 文档，逐页提取文本、表格、图片描述和目录"""
        logger.info(f"解析 PDF: {file_path}")
        doc = fitz.open(file_path)

        pages = []
        for i, page in enumerate(doc):
            page_parts: list[str] = []
            page_meta = {
                "width": page.rect.width,
                "height": page.rect.height,
                "table_count": 0,
                "image_count": 0,
                "word_count": 0,
            }

            # ---- 1. 提取纯文本（按自然阅读顺序）----
            text_blocks = page.get_text("text").strip()
            if text_blocks:
                page_parts.append(text_blocks)
                page_meta["word_count"] = len(text_blocks.split())

            # ---- 2. 提取表格（PyMuPDF 1.23+ 支持）----
            tables_text = _extract_pdf_tables(page, i + 1)
            if tables_text:
                page_parts.append(tables_text)
                page_meta["table_count"] = tables_text.count("【表格")

            # ---- 3. 提取图片描述（图片位置、尺寸、标题线索）----
            images_text = _extract_pdf_images(page, i + 1, doc)
            if images_text:
                page_parts.append(images_text)
                page_meta["image_count"] = images_text.count("【图片")

            # ---- 4. 合并页面内容 ----
            combined_text = "\n\n".join(p for p in page_parts if p.strip())

            pages.append(ParsedPage(
                page_num=i + 1,
                text=combined_text,
                metadata=page_meta,
            ))

        # 提取 PDF 内置目录
        toc = []
        for item in doc.get_toc():
            toc.append({
                "level": item[0],
                "title": item[1],
                "page": item[2]
            })

        title = doc.metadata.get("title", Path(file_path).stem)
        doc.close()

        return ParsedDocument(
            file_path=file_path,
            title=title,
            pages=pages,
            raw_text="\n".join(p.text for p in pages),
            toc=toc
        )

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """解析 DOCX 文档"""
        logger.info(f"解析 DOCX: {file_path}")
        doc = DocxDocument(file_path)

        full_text_parts = []
        current_page = 1
        page_texts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检查段落样式判断是否为标题（用于构建目录）
            if para.style.name.startswith("Heading"):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                page_texts.append((current_page, text, {"is_heading": True, "level": level}))
            else:
                page_texts.append((current_page, text, {"is_heading": False}))

            full_text_parts.append(text)

        # DOCX 没有分页信息，将所有内容视为一页
        pages = [ParsedPage(
            page_num=1,
            text="\n".join(full_text_parts),
            metadata={"paragraphs": len(full_text_parts)}
        )]

        return ParsedDocument(
            file_path=file_path,
            title=Path(file_path).stem,
            pages=pages,
            raw_text="\n".join(full_text_parts),
            toc=[]
        )

    def _parse_text(self, file_path: str) -> ParsedDocument:
        """解析纯文本 / Markdown 文档"""
        logger.info(f"解析文本: {file_path}")
        raw_text = Path(file_path).read_text(encoding="utf-8")

        return ParsedDocument(
            file_path=file_path,
            title=Path(file_path).stem,
            pages=[ParsedPage(page_num=1, text=raw_text)],
            raw_text=raw_text,
            toc=[]
        )

    def _parse_excel(self, file_path: str) -> ParsedDocument:
        """
        解析 Excel 文档，将每个 Sheet 作为一页，
        每行转为制表符分隔的文本（后续由 Markdown 转换器转为表格）。
        """
        logger.info(f"解析 Excel: {file_path}")
        path = Path(file_path)
        ext = path.suffix.lower()

        pages = []

        if ext == ".xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows_text = []
                for row in ws.iter_rows(values_only=True):
                    # 将行转为 tab 分隔的文本，便于后续转为 Markdown 表格
                    row_str = "\t".join(
                        str(cell) if cell is not None else ""
                        for cell in row
                    )
                    if row_str.strip():
                        rows_text.append(row_str)

                sheet_text = f"【工作表: {sheet_name}】\n" + "\n".join(rows_text)
                pages.append(ParsedPage(
                    page_num=len(pages) + 1,
                    text=sheet_text,
                    metadata={"sheet_name": sheet_name, "rows": len(rows_text)}
                ))
            wb.close()

        elif ext == ".xls":
            import xlrd
            wb = xlrd.open_workbook(file_path)
            for sheet in wb.sheets():
                rows_text = []
                for row_idx in range(sheet.nrows):
                    row_values = sheet.row_values(row_idx)
                    row_str = "\t".join(
                        str(cell) if cell != "" else ""
                        for cell in row_values
                    )
                    if row_str.strip():
                        rows_text.append(row_str)

                sheet_text = f"【工作表: {sheet.name}】\n" + "\n".join(rows_text)
                pages.append(ParsedPage(
                    page_num=len(pages) + 1,
                    text=sheet_text,
                    metadata={"sheet_name": sheet.name, "rows": len(rows_text)}
                ))

        raw_text = "\n\n".join(p.text for p in pages)

        return ParsedDocument(
            file_path=file_path,
            title=path.stem,
            pages=pages,
            raw_text=raw_text,
            toc=[]
        )

    def _parse_image(self, file_path: str) -> ParsedDocument:
        """
        解析图片：提取图像元信息，可选 OCR。
        图片的文本提取主要由 Kimi Vision 或 OCR 完成后，
        作为 Markdown 清洗流程的一部分。
        """
        logger.info(f"解析图片: {file_path}")
        path = Path(file_path)

        try:
            from PIL import Image
            img = Image.open(file_path)
            width, height = img.size
            img_format = img.format or path.suffix.upper().lstrip(".")
            img_mode = img.mode

            # 尝试 OCR 提取文字
            ocr_text = self._ocr_image(file_path)

            # 构建基础描述文本
            desc_lines = [
                f"# 图片: {path.name}",
                f"",
                f"- 尺寸: {width}x{height}",
                f"- 格式: {img_format}",
                f"- 色彩模式: {img_mode}",
                f"",
            ]
            if ocr_text.strip():
                desc_lines.append(f"## OCR 识别文字\n\n{ocr_text}")
            else:
                desc_lines.append(f"> 此图片未检测到可识别的文字，可能需要视觉模型提取内容。")

            raw_text = "\n".join(desc_lines)

            return ParsedDocument(
                file_path=file_path,
                title=path.stem,
                pages=[ParsedPage(
                    page_num=1,
                    text=raw_text,
                    metadata={
                        "width": width,
                        "height": height,
                        "format": img_format,
                        "mode": img_mode,
                    }
                )],
                raw_text=raw_text,
                toc=[]
            )

        except ImportError:
            logger.warning("Pillow 未安装，使用最简图片解析")
            raw_text = f"# 图片: {Path(file_path).name}\n\n(图片内容需通过视觉模型提取)"
            return ParsedDocument(
                file_path=file_path,
                title=Path(file_path).stem,
                pages=[ParsedPage(page_num=1, text=raw_text)],
                raw_text=raw_text,
                toc=[]
            )

    def _ocr_image(self, file_path: str) -> str:
        """尝试对图片执行 OCR 文字识别"""
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            return text.strip()
        except ImportError:
            logger.debug("pytesseract 未安装，跳过 OCR")
            return ""
        except Exception as e:
            logger.warning(f"OCR 识别失败: {e}")
            return ""


# =================================================================
# PDF 增强解析：表格 + 图片描述 + 结构化文本
# =================================================================

def _extract_pdf_tables(page, page_num: int) -> str:
    """
    从 PDF 页面提取表格，转换为 Markdown 格式。

    使用 PyMuPDF 1.23+ 的 page.find_tables() 接口。
    每个表格输出格式：
        【表格 N · 第 page_num 页】
        | 列 1 | 列 2 | ...
        |------|------| ...
        | 单元格内容 | ...
    """
    try:
        tables = page.find_tables()
    except Exception as e:
        # 旧版 PyMuPDF 不支持 find_tables，降级为忽略
        if "'Page' object has no attribute 'find_tables" in str(e):
            logger.debug(f"PDF 页 {page_num}: 当前 PyMuPDF 版本不支持 find_tables，跳过表格提取")
        else:
            logger.debug(f"PDF 页 {page_num}: 表格提取失败: {e}")
        return ""

    if not tables or not tables.tables:
        return ""

    output_parts: list[str] = []
    for t_idx, table in enumerate(tables.tables, 1):
        try:
            raw_rows = table.extract()
        except Exception:
            continue

        if not raw_rows:
            continue

        # 过滤完全空行
        rows = []
        for r in raw_rows:
            if any((c or "").strip() for c in r):
                rows.append([(c or "").strip().replace("\n", " ") for c in r])

        if not rows:
            continue

        header = rows[0]
        body = rows[1:]

        # 生成 Markdown 表格
        col_count = max(len(header), max((len(r) for r in body), default=0))
        if col_count == 0:
            continue

        md_lines = [f"【表格 {t_idx} · 第 {page_num} 页】"]
        md_lines.append("| " + " | ".join((header + [""] * col_count)[:col_count]) + " |")
        md_lines.append("| " + " | ".join(["---"] * col_count) + " |")
        for row in body[:50]:  # 每个表格最多保留 50 行
            md_lines.append("| " + " | ".join((row + [""] * col_count)[:col_count]) + " |")
        if len(body) > 50:
            md_lines.append(f"| ...（共 {len(body)} 行，已截断） |")
        output_parts.append("\n".join(md_lines))

    return "\n\n".join(output_parts) if output_parts else ""


def _extract_pdf_images(page, page_num: int, doc) -> str:
    """
    从 PDF 页面提取图片信息，生成语义描述文本。

    输出格式：
        【图片 N · 第 page_num 页】
        - 位置: 左上方 / 居中 / 右下方
        - 尺寸: 宽度 × 高度 像素
        - 相邻文本: "图片下方或上方的上下文文字"
    """
    try:
        image_list = page.get_images(full=True)
    except Exception as e:
        logger.debug(f"PDF 页 {page_num}: 图片提取失败: {e}")
        return ""

    if not image_list:
        return ""

    page_rect = page.rect
    page_width = max(page_rect.width, 1)
    page_height = max(page_rect.height, 1)

    # 提取该页文本块（按布局），用于推断图片上下文
    try:
        text_dict = page.get_text("dict")
        text_blocks = [
            {
                "text": " ".join(
                    span.get("text", "")
                    for line in block.get("lines", [])
                    for span in line.get("spans", [])
                ).strip(),
                "bbox": block.get("bbox", (0, 0, 0, 0)),
            }
            for block in text_dict.get("blocks", [])
            if block.get("type") == 0  # 0 = text block
        ]
    except Exception:
        text_blocks = []

    output_parts: list[str] = []
    for idx, img in enumerate(image_list[:10], 1):  # 每页最多记录 10 张图片
        try:
            xref = img[0]
            # 获取图片在页面上的位置
            try:
                bbox_list = page.get_image_rects(xref)
                if not bbox_list:
                    continue
                bbox = bbox_list[0]
            except Exception:
                continue

            x0, y0, x1, y1 = bbox
            img_w = max(int(x1 - x0), 1)
            img_h = max(int(y1 - y0), 1)

            # 判断图片在页面上的位置
            cx = (x0 + x1) / 2.0
            cy = (y0 + y1) / 2.0
            position_h = "左侧" if cx < page_width / 3 else ("右侧" if cx > page_width * 2 / 3 else "居中")
            position_v = "上方" if cy < page_height / 3 else ("下方" if cy > page_height * 2 / 3 else "中部")
            position = f"{position_v}{position_h}"

            # 找最近的文本块（作为图片上下文线索）
            adjacent_text = ""
            best_dist = float("inf")
            for tb in text_blocks:
                if not tb["text"]:
                    continue
                bx0, by0, bx1, by1 = tb["bbox"]
                cx2, cy2 = (bx0 + bx1) / 2.0, (by0 + by1) / 2.0
                dist = abs(cx2 - cx) + abs(cy2 - cy)
                if dist < best_dist:
                    best_dist = dist
                    adjacent_text = tb["text"]

            # 从 PDF 对象中尝试获取图片描述元信息
            img_size_hint = ""
            try:
                img_obj = doc.extract_image(xref)
                if img_obj:
                    w = img_obj.get("width")
                    h = img_obj.get("height")
                    if w and h:
                        img_size_hint = f"{w}×{h}"
            except Exception:
                pass

            desc_lines = [f"【图片 {idx} · 第 {page_num} 页】"]
            desc_lines.append(f"- 位置: {position}")
            desc_lines.append(f"- 尺寸: {img_w}×{img_h} 页面单位" + (f"（原始 {img_size_hint} 像素）" if img_size_hint else ""))
            if adjacent_text:
                snippet = adjacent_text[:80]
                desc_lines.append(f"- 相邻文本: \"{snippet}{'...' if len(adjacent_text) > 80 else ''}\"")
            desc_lines.append(f"- 类型: 嵌入图片（若需识别文字，请启用 OCR）")

            output_parts.append("\n".join(desc_lines))
        except Exception as e:
            logger.debug(f"PDF 页 {page_num} 图片 {idx} 解析异常: {e}")
            continue

    return "\n\n".join(output_parts) if output_parts else ""